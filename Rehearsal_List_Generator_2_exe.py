#! /usr/bin/env python
#  -*- coding: utf-8 -*-
#
# GUI module generated by PAGE version 4.20
#  in conjunction with Tcl version 8.6
#    Apr 30, 2019 01:38:11 PM MDT  platform: Windows NT
#By Bart Cubrich
import sys

try:
    from Tkinter import *
    import Tkinter as tk
except ImportError:
    from tkinter import *
    import tkinter as tk

from tkinter.filedialog import askopenfilename
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox


try:
    import ttk
    py3 = False
except ImportError:
    import tkinter.ttk as ttk
    py3 = True

import Rehearsal_List_Generator_2_support

import pandas as pd
import numpy as np
from numpy.random import choice
from docx import Document
from docx.shared import Inches
import os
import psutil



def vp_start_gui():
    '''Starting point when module is the main routine.'''
    global val, w, root
    root = tk.Tk()
    top = Toplevel1 (root)
    Rehearsal_List_Generator_2_support.init(root, top)
    root.mainloop()

w = None
def create_Toplevel1(root, *args, **kwargs):
    '''Starting point when module is imported by another program.'''
    global w, w_win, rt
    rt = root
    w = tk.Toplevel (root)
    top = Toplevel1 (w)
    Rehearsal_List_Generator_2_support.init(w, top, *args, **kwargs)
    return (w, top)

def destroy_Toplevel1():
    global w
    w.destroy()
    w = None

class Toplevel1:

    def __init__(self, top=None):
        self.book_file=None
        self.save_file=None
        '''This class configures and populates the toplevel window.
           top is the toplevel containing window.'''
        _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
        _fgcolor = '#000000'  # X11 color: 'black'
        _compcolor = '#d9d9d9' # X11 color: 'gray85'
        _ana1color = '#d9d9d9' # X11 color: 'gray85' 
        _ana2color = '#ececec' # Closest X11 color: 'gray92' 
        font9 = "-family {Segoe UI} -size 9 -weight normal -slant "  \
            "roman -underline 0 -overstrike 0"
        self.style = ttk.Style()
        if sys.platform == "win32":
            self.style.theme_use('winnative')
        self.style.configure('.',background=_bgcolor)
        self.style.configure('.',foreground=_fgcolor)
        self.style.configure('.',font="TkDefaultFont")
        self.style.map('.',background=
            [('selected', _compcolor), ('active',_ana2color)])

        top.geometry("600x450")
        top.title("Rehearsal List Generator")
        top.configure(background="#d9d9d9")
        
                      
                      
                      
        self.instruction_set='''To use this program follow the steps below
        1. Choose the location of the file with the songs info and probabilities
          a.  probabilities are simply any numbers showing how often you want to 
                see the song
          b.  in the simplest case you leave the probabilities blank 
                everywhere except for on song you want more often
                i. in this case the number you choose is how much 
                    more likely that song is to be chosen. In other
                    words, all weights are relative to one if only
                    positive numbers are used.
                ii. e.g, a 2 is twice as likely as a 1
                iii. using numbers close to 1000 for a single 
                      song will make that song almost 100% likely 
                      for small books (~200 songs)
                iv. negative probabilities will make a song less likely, 
                    but will make positive number harder to interpret
                v. intergers are fine to use as weights, and there is 
                   no good reason to use decimals
         
        2. Choose where you want to save the file by 
            specify a file name (.doc or .docx)
        
        3. Enter the rehearsal data the list is for
        
        4. Choose whether you want christmas charts or not
        
        5. Choose whether you want vocal charts or not
        
        6. Press the generate button
          a. the generated list will be shown as a popup
          b. if you don't like it but want the same setting you 
               can just keep hitting generate over and over'''
               
        self.book_file_name=tk.StringVar(top, value='')
        self.Text1 = tk.Entry(top,textvariable=self.book_file_name)
        self.Text1.place(relx=0.033, rely=0.133, relheight=0.053, relwidth=0.75)

        self.Text1.configure(background="#C9C9C9")
        self.Text1.configure(font=font9)
        self.Text1.configure(foreground="black")
        self.Text1.configure(highlightbackground="#d9d9d9")
        self.Text1.configure(highlightcolor="black")
        self.Text1.configure(insertbackground="black")
        self.Text1.configure(selectbackground="#C9C9C9")
        self.Text1.configure(selectforeground="black")
        self.Text1.configure(width=314)
#        self.Text1.configure(wrap='word')

        self.Button1 = tk.Button(top,command= self.get_book_file)
        self.Button1.place(relx=0.8, rely=0.133, height=24, width=120)
        self.Button1.configure(activebackground="#ececec")
        self.Button1.configure(activeforeground="#000000")
        self.Button1.configure(background="#d8d582")
        self.Button1.configure(disabledforeground="#a3a3a3")
        self.Button1.configure(foreground="#000000")
        self.Button1.configure(highlightbackground="#d9d9d9")
        self.Button1.configure(highlightcolor="black")
        self.Button1.configure(pady="0")
        self.Button1.configure(text='''Choose Songbook''')
        self.Button1.configure(width=187)


        self.save_file_name=tk.StringVar(top, value='')
        self.Text2 = tk.Entry(top,textvariable=self.save_file_name)
        self.Text2.place(relx=0.033, rely=0.267, relheight=0.053, relwidth=0.7)

        self.Text2.configure(background="#C9C9C9")
        self.Text2.configure(font=font9)
        self.Text2.configure(foreground="black")
        self.Text2.configure(highlightbackground="#d9d9d9")
        self.Text2.configure(highlightcolor="black")
        self.Text2.configure(insertbackground="black")
        self.Text2.configure(selectbackground="#c4c4c4")
        self.Text2.configure(selectforeground="black")
        self.Text2.configure(width=314)
#        self.Text2.configure(wrap='word')
        
        self.rehearsal_date=tk.StringVar(top, value='Enter Rehearsal Date')
        self.Date_Box = tk.Entry(top,textvariable=self.rehearsal_date)
        self.Date_Box.place(relx=0.033, rely=0.35, relheight=0.053, relwidth=0.2)        
         

        self.Button2 = tk.Button(top,command=self.get_save_file)
        self.Button2.place(relx=0.75, rely=0.267, height=24, width=130)
        self.Button2.configure(activebackground="#ececec")
        self.Button2.configure(activeforeground="#000000")
        self.Button2.configure(background="#d8d582")
        self.Button2.configure(disabledforeground="#a3a3a3")
        self.Button2.configure(foreground="#000000")
        self.Button2.configure(highlightbackground="#d9d9d9")
        self.Button2.configure(highlightcolor="black")
        self.Button2.configure(pady="0")
        self.Button2.configure(text='''Choose Save Location''')
        self.Button2.configure(width=187)

        self.Label1 = tk.Label(top)
        self.Label1.place(relx=0.033, rely=0.067, height=21, width=150)
        self.Label1.configure(background="#d9d9d9")
        self.Label1.configure(disabledforeground="#a3a3a3")
        self.Label1.configure(foreground="#000000")
        self.Label1.configure(text='''Choose Songbook Excel File''')

        self.Label2 = tk.Label(top)
        self.Label2.place(relx=0.033, rely=0.2, height=21, width=225)
        self.Label2.configure(background="#d9d9d9")
        self.Label2.configure(disabledforeground="#a3a3a3")
        self.Label2.configure(foreground="#000000")
        self.Label2.configure(text='''Choose Output MS Word Save As Filename''')

        self.rad1_var = IntVar(top)
        
#        self.style.map('TRadiobutton',background=
#            [('selected', _bgcolor), ('active', _ana2color)])
        common_fg='#000000'
        common_bg='#d9d9d9'
        self.TRadiobutton1 = tk.Radiobutton(top,variable=self.rad1_var, value=1, fg=common_fg, bg=common_bg)
        self.TRadiobutton1.place(relx=0.05, rely=0.444, relwidth=0.217
                , relheight=0.0, height=21)
        self.TRadiobutton1.configure(takefocus="")
        self.TRadiobutton1.configure(text='''No Christmas Music''')
#        self.TRadiobutton1.configure(activebackground="#d9d9d9")

        self.TRadiobutton2 = tk.Radiobutton(top,variable=self.rad1_var, value=2, fg=common_fg, bg=common_bg)
        self.TRadiobutton2.place(relx=0.283, rely=0.444, relwidth=0.232
                , relheight=0.0, height=21)
        self.TRadiobutton2.configure(takefocus="")
        self.TRadiobutton2.configure(text='''Only Chirstmas Music''')
        
        self.TRadiobutton3 = tk.Radiobutton(top,variable=self.rad1_var, value=3, fg=common_fg, bg=common_bg)
        self.TRadiobutton3.place(relx=0.533, rely=0.444, relwidth=0.262
                , relheight=0.0, height=21)
        self.TRadiobutton3.configure(takefocus="")
        self.TRadiobutton3.configure(text='''Christmas Music Possible''')
        self.rad1_var.set(1)
#        self.TRadiobutton1.select()
        
        self.rad2_var = IntVar(top)
        
        self.TRadiobutton4 = tk.Radiobutton(top,variable=self.rad2_var, value=1, fg=common_fg, bg=common_bg)
        self.TRadiobutton4.place(relx=0.05, rely=0.556, relwidth=0.18
                , relheight=0.0, height=21)
        self.TRadiobutton4.configure(takefocus="")
        self.TRadiobutton4.configure(text='''No Vocal Charts''')

        self.TRadiobutton5 = tk.Radiobutton(top,variable=self.rad2_var, value=2, fg=common_fg, bg=common_bg)
        self.TRadiobutton5.place(relx=0.283, rely=0.556, relwidth=0.195
                , relheight=0.0, height=21)
        self.TRadiobutton5.configure(takefocus="")
        self.TRadiobutton5.configure(text='''Only Vocal Charts''')

        self.TRadiobutton6 = tk.Radiobutton(top,variable=self.rad2_var, value=3, fg=common_fg, bg=common_bg)
        self.TRadiobutton6.place(relx=0.533, rely=0.556, relwidth=0.225
                , relheight=0.0, height=21)
        self.TRadiobutton6.configure(takefocus="")
        self.TRadiobutton6.configure(text='''Vocal Charts Possible''')
        self.rad2_var.set(3)
#        self.TRadiobutton6.select()
        
        self.Button3 = tk.Button(top, command=self.generate)
        self.Button3.place(relx=0.033, rely=0.8, height=74, width=157)
        self.Button3.configure(activebackground="#ececec")
        self.Button3.configure(activeforeground="#000000")
        self.Button3.configure(background="#63d85f")
        self.Button3.configure(disabledforeground="#a3a3a3")
        self.Button3.configure(foreground="#000000")
        self.Button3.configure(highlightbackground="#d9d9d9")
        self.Button3.configure(highlightcolor="black")
        self.Button3.configure(pady="0")
        self.Button3.configure(text='''Generate Rehearsal List''')
        self.Button3.configure(width=157)
        

        
        self.Button4 = tk.Button(top,command=self.quitit)
        self.Button4.place(relx=0.85, rely=0.889, height=34, width=84)
        self.Button4.configure(activebackground="#ececec")
        self.Button4.configure(activeforeground="#000000")
        self.Button4.configure(background="#d83838")
        self.Button4.configure(disabledforeground="#a3a3a3")
        self.Button4.configure(foreground="#000000")
        self.Button4.configure(highlightbackground="#d9d9d9")
        self.Button4.configure(highlightcolor="black")
        self.Button4.configure(pady="0")
        self.Button4.configure(text='''Quit''')
        self.Button4.configure(width=84)
        
        self.Button5 = tk.Button(top,command=self.instructions)
        self.Button5.place(relx=0.5, rely=0.889, height=35, width=85)
        self.Button5.configure(activebackground="#ececec")
        self.Button5.configure(activeforeground="#000000")
        self.Button5.configure(background="#d8d582")
        self.Button5.configure(disabledforeground="#a3a3a3")
        self.Button5.configure(foreground="#000000")
        self.Button5.configure(highlightbackground="#d8d582")
        self.Button5.configure(highlightcolor="black")
        self.Button5.configure(pady="0")
        self.Button5.configure(text='''Instructions''')
        self.Button5.configure(width=85)
        
        self.Label3 = tk.Label(top)
        self.Label3.place(relx=0.033, rely=0.65, height=21, width=200)
        self.Label3.configure(background="#d9d9d9")
        self.Label3.configure(disabledforeground="#a3a3a3")
        self.Label3.configure(foreground="#000000")
        self.Label3.configure(text='''Choose Number of Songs to Rehearse''')
        
        self.n_songs=tk.IntVar(top, value=15)
        self.Text1 = tk.Entry(top,textvariable=self.n_songs)
        self.Text1.place(relx=0.033, rely=0.7, relheight=0.053, relwidth=0.1)
        
    def quitit(self):
        
        root.quit()
        root.destroy()
        sys.exit("Script No Longer Running")
        
    def get_book_file(self):
        self.book_file = askopenfilename(title = "Select Song Book File With Probabilities")      # Open single file
        self.book_file_name.set(self.book_file) 
    
    def get_save_file(self):
        self.save_file = asksaveasfilename(title = "Save As", filetypes=[('.docx', '.docx'),('.doc', '.doc')],defaultextension='.docx')      # Open single file
        self.save_file_name.set(self.save_file ) 
    
    def instructions(self):
        messagebox.showinfo("Instructions", self.instruction_set)
        print(self.rad1_var.get())
        print(self.rad2_var.get())
    
    def generate(self):
        
        if self.rad1_var.get()==2 and self.rad2_var.get()==2:
            messagebox.showinfo("Selection Error", "Can't choose only christmas and only vocal at the same time. Please make another selection.")
        else:
            if self.book_file!=None:
                data_file=self.book_file
                
                df=pd.read_excel(data_file, index_col =0)
                df.index=df.index.astype(str)
                
                df['Title']=df['Title'].replace(' ',np.NaN)
                df=df.dropna(subset=['Title'])
                df=df[df.index.notnull()]
                
                #check if the user want christmas songs
                if self.rad1_var.get()==1:
                    df=df[~df.index.str.contains('C')]
                elif self.rad1_var.get()==2:
                    df=df[df.index.str.contains('C')]
                
                #check if the user wants vocal songs
                if self.rad2_var.get()==1:
                    df=df[~df.index.str.contains('V')]
                elif self.rad2_var.get()==2:
                    df=df[df.index.str.contains('V')]
                
                #fil in ones where the user didn't give a weight
                df['Extra Weight']=df['Extra Weight'].fillna(1)
                
                if any(df['Extra Weight']<0):
                    df['Extra Weight']=np.where(df['Extra Weight']<0,(df['Extra Weight']+np.abs(df['Extra Weight'].min())+1)/np.abs(df['Extra Weight']),df['Extra Weight']*df['Extra Weight']+np.abs(df['Extra Weight'])+1)
                  
                
                
                df['Probability']=df['Extra Weight']/df['Extra Weight'].sum()
                
                
                draw = list(choice(df.index, self.n_songs.get(), p=df['Probability'], replace=False))
                
                
                sample=df.loc[draw].drop(['Extra Weight','Probability'], axis='columns')
                sample=sample.reset_index()
                ordered_cols=['Number','Title', 'Tempo','Minutes','Seconds']
                sample=sample[ordered_cols].fillna('')
                
                out_table=sample.copy()[['Number','Title']].to_string()
                
    
#                document = Document(docx='default.docx')
                document = Document()
                user_date=self.rehearsal_date.get()
                
                document.add_heading('Riverton Jazz Band Rehearsal List {}'.format(user_date), 0)
                
                t = document.add_table(sample.shape[0]+1, sample.shape[1])
                
                # add the header rows.
                for j in range(sample.shape[-1]):
                    t.cell(0,j).text = sample.columns[j]
                
                # add the rest of the data frame
                for i in range(sample.shape[0]):
                    for j in range(sample.shape[-1]):
                        t.cell(i+1,j).text = str(sample.values[i,j])
                
                def set_col_widths(table):
                    widths = (Inches(0.5),Inches(3), Inches(0.5), Inches(0.5),Inches(0.5))
                    for row in table.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                            
                set_col_widths(t)
                
                
                
                if self.save_file != None:
                    messagebox.showinfo("Selected songs", out_table)
                    out_file=self.save_file
                    if out_file.endswith('.docx')==False or out_file.endswith('.doc')==False:
                        out_file=out_file.split('.')[0]+'.docx'
                    try:
                        document.save(out_file)
                    except PermissionError:
                        messagebox.showerror("Permission Error", "Permission Denied. Please Close Word Document Before Running")
                else:
                    messagebox.showerror("Song Book File Error", "No save file name. Please choose a a save file.")
            else:
                messagebox.showerror("Song Book File Error", "No song book file selected. Please choose an excel file.")

if __name__ == '__main__':
    vp_start_gui()





